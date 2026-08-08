from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

NAVY = "17365D"
BLUE = "2563EB"
TEAL = "0F766E"
PURPLE = "7C3AED"
ORANGE = "EA580C"
GREEN = "15803D"
RED = "B91C1C"
INK = "1F2937"
MUTED = "64748B"
LIGHT = "F8FAFC"
WHITE = "FFFFFF"


def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rounded(draw, box, fill, outline=None, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill="#" + fill, outline=("#" + outline if outline else None), width=width)


def center_text(draw, box, text, size=28, color=INK, bold=False, line_gap=7):
    f = font(size, bold)
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=f)[3] for line in lines]
    total = sum(heights) + line_gap * (len(lines) - 1)
    y = box[1] + (box[3] - box[1] - total) / 2
    for line, h in zip(lines, heights):
        bb = draw.textbbox((0, 0), line, font=f)
        x = box[0] + (box[2] - box[0] - (bb[2] - bb[0])) / 2
        draw.text((x, y), line, font=f, fill="#" + color)
        y += h + line_gap


def arrow(draw, start, end, color=MUTED, width=6):
    draw.line([start, end], fill="#" + color, width=width)
    x, y = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        direction = 1 if end[0] > start[0] else -1
        pts = [(x, y), (x - 18 * direction, y - 12), (x - 18 * direction, y + 12)]
    else:
        direction = 1 if end[1] > start[1] else -1
        pts = [(x, y), (x - 12, y - 18 * direction), (x + 12, y - 18 * direction)]
    draw.polygon(pts, fill="#" + color)


def workflow_image(path):
    img = Image.new("RGB", (2400, 1500), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((100, 65), "AI UX Discovery-to-Prototype Workflow", font=font(58, True), fill="#" + NAVY)
    d.text((104, 135), "Agent-owned execution from BRD to a human-approved synthetic demo prototype", font=font(28), fill="#" + MUTED)

    colors = [BLUE, TEAL, PURPLE, ORANGE]
    phases = [
        ("PHASE 1", "Discover", "UX Research Agent\nRequirements Challenge Agent", "Evidence analysis\nBRD challenge\nPersonas & pain points", "D1"),
        ("PHASE 2", "Define & Ideate", "UX Definition &\nIdeation Agent", "Problem definition\nJourney & opportunities\nConcept selection", "D2"),
        ("PHASE 3", "Design & Prototype", "Experience Design Agent", "User flow & screens\nInteraction states\nRunnable prototype", "D3"),
        ("PHASE 4", "Validate", "UX Validation &\nAudit Agent", "Synthetic walkthrough\nAccessibility & coverage\nValidation pack", "D4"),
    ]
    x0, y0, w, gap = 110, 330, 460, 95
    for i, (tag, title, agents, body, gate) in enumerate(phases):
        x = x0 + i * (w + gap)
        rounded(d, (x, y0, x + w, y0 + 525), WHITE, colors[i], 32, 5)
        rounded(d, (x + 28, y0 + 28, x + 165, y0 + 82), colors[i], None, 18)
        center_text(d, (x + 28, y0 + 28, x + 165, y0 + 82), tag, 21, WHITE, True)
        center_text(d, (x + 25, y0 + 100, x + w - 25, y0 + 175), title, 33, colors[i], True)
        agent_box = (x + 28, y0 + 185, x + w - 28, y0 + 290)
        rounded(d, agent_box, colors[i], None, 18)
        center_text(d, agent_box, agents, 22, WHITE, True, 7)
        center_text(d, (x + 35, y0 + 305, x + w - 35, y0 + 430), body, 24, INK, False, 10)
        gate_box = (x + 125, y0 + 445, x + w - 125, y0 + 505)
        rounded(d, gate_box, colors[i], None, 20)
        center_text(d, gate_box, f"HUMAN GATE {gate}", 20, WHITE, True)
        if i < 3:
            arrow(d, (x + w + 12, y0 + 265), (x + w + gap - 12, y0 + 265), colors[i], 7)

    rounded(d, (690, 1030, 1710, 1215), "ECFDF5", GREEN, 34, 5)
    center_text(d, (710, 1045, 1690, 1200), "VERIFIED WORKABLE PROTOTYPE\nfor the documented synthetic demo scope", 35, GREEN, True, 10)
    arrow(d, (2140, 855), (1710, 1120), GREEN, 8)

    rounded(d, (110, 1300, 2290, 1415), "FFF7ED", ORANGE, 24, 3)
    center_text(d, (135, 1310, 2265, 1405), "APPROVE advances  •  REVISE returns only to affected work  •  REJECT stops the workflow  •  Agents never self-approve", 25, INK, True)
    img.save(path, quality=95)


def folder_image(path):
    img = Image.new("RGB", (2400, 1650), "#F8FAFC")
    d = ImageDraw.Draw(img)

    # Dark presentation header matching the supplied visual reference.
    d.rectangle((0, 0, 2400, 150), fill="#073A72")
    d.rectangle((45, 72, 310, 82), fill="#18C4C7")
    d.rectangle((2090, 72, 2355, 82), fill="#18C4C7")
    center_text(d, (320, 18, 2080, 132), "AI UX PIPELINE — PROJECT FOLDER STRUCTURE", 49, WHITE, True)

    def folder_icon(x, y, scale=1.0):
        d.rounded_rectangle((x, y + 9*scale, x + 46*scale, y + 39*scale), radius=5*scale, fill="#FFC83D", outline="#B27600", width=max(1, int(2*scale)))
        d.rounded_rectangle((x + 4*scale, y, x + 23*scale, y + 14*scale), radius=4*scale, fill="#FFD760", outline="#B27600", width=max(1, int(2*scale)))

    def file_icon(x, y, scale=1.0):
        d.rectangle((x, y, x + 30*scale, y + 38*scale), fill="#FFFFFF", outline="#2B65A8", width=max(1, int(2*scale)))
        d.polygon([(x + 20*scale, y), (x + 30*scale, y + 10*scale), (x + 20*scale, y + 10*scale)], fill="#D7E8FA")
        for yy in (17, 24, 31):
            d.line((x + 6*scale, y + yy*scale, x + 24*scale, y + yy*scale), fill="#75A1CF", width=max(1, int(scale)))

    def panel(box, title, rows, accent=BLUE, title_icon="folder", text_size=19, row_gap=37):
        x1, y1, x2, y2 = box
        rounded(d, box, "F2F8FF", "A7CFF2", 18, 3)
        if title_icon == "folder": folder_icon(x1 + 24, y1 + 17, .9)
        d.text((x1 + 86, y1 + 23), title, font=font(25, True), fill="#" + NAVY)
        branch_x = x1 + 70
        start_y = y1 + 75
        d.line((branch_x, start_y - 15, branch_x, start_y + (len(rows)-1)*row_gap + 15), fill="#" + NAVY, width=3)
        for i, (kind, label) in enumerate(rows):
            yy = start_y + i * row_gap
            d.line((branch_x, yy + 12, branch_x + 38, yy + 12), fill="#" + NAVY, width=3)
            if kind == "folder": folder_icon(branch_x + 43, yy - 5, .65)
                
            else: file_icon(branch_x + 50, yy - 7, .62)
            d.text((branch_x + 92, yy - 4), label, font=font(text_size, kind == "folder"), fill="#" + INK)

    # Root container and repository root.
    rounded(d, (60, 175, 2340, 1435), WHITE, NAVY, 22, 4)
    folder_icon(120, 205, 1.25)
    d.text((190, 218), "ai-ux-design-pipeline/", font=font(31, True), fill="#" + NAVY)
    d.line((155, 260, 155, 1350), fill="#" + NAVY, width=4)

    # Left column: workflow, agents, skills. Agent filenames are intentionally prominent.
    panel((155, 275, 1155, 540), "workflow/", [
        ("file", "orchestrator.md"),
        ("file", "run-phase-01-discover.md"),
        ("file", "run-phase-02-define.md"),
        ("file", "run-phase-03-design.md"),
        ("file", "run-phase-04-validate.md"),
    ], BLUE, text_size=18, row_gap=36)
    panel((155, 560, 1155, 870), "agents/  —  RESPONSIBILITY OWNERS", [
        ("file", "requirements-challenge-agent.md"),
        ("file", "ux-research-agent.md"),
        ("file", "ux-definition-agent.md"),
        ("file", "experience-design-agent.md"),
        ("file", "ux-validation-audit-agent.md"),
    ], TEAL, text_size=20, row_gap=44)
    panel((155, 890, 1155, 1330), "skills/  —  22 REUSABLE CAPABILITIES", [
        ("folder", "Requirements challenge: 3 skills"),
        ("folder", "Research synthesis: 3 skills"),
        ("folder", "Definition & ideation: 5 skills"),
        ("folder", "Design & prototype: 5 skills"),
        ("folder", "Validation & audit: 6 skills"),
    ], PURPLE, text_size=19, row_gap=57)

    # Right column: gates, inputs, outputs, runtime/tests.
    panel((1245, 275, 2245, 505), "gates/  —  HUMAN-IN-THE-LOOP", [
        ("file", "gate-d1-research-review.md"),
        ("file", "gate-d2-definition-review.md"),
        ("file", "gate-d3-prototype-review.md"),
        ("file", "gate-d4-final-validation-review.md"),
    ], ORANGE, text_size=18, row_gap=36)
    panel((1245, 525, 2245, 735), "projects/starter/input/", [
        ("file", "sample-brd.md + stakeholder-notes.md"),
        ("file", "existing-flow.md + design-system.md"),
        ("file", "4 synthetic research input files"),
    ], TEAL, text_size=18, row_gap=42)
    panel((1245, 755, 2245, 1095), "outputs/", [
        ("folder", "phase-01/  research + BRD challenge"),
        ("folder", "phase-02/  definition + ideation"),
        ("folder", "phase-03/  design + prototype/"),
        ("folder", "phase-04/  validation + audit"),
        ("folder", "reviews/  d1-review.md … d4-review.md"),
    ], GREEN, text_size=18, row_gap=47)
    panel((1245, 1115, 1725, 1330), "src/ai_ux_workflow/", [
        ("file", "runtime.py + cli.py"),
        ("file", "contracts.py + challenge.py"),
    ], BLUE, text_size=16, row_gap=44)
    panel((1765, 1115, 2245, 1330), "tests/", [
        ("file", "6 test modules"),
        ("file", "32 passing tests"),
    ], PURPLE, text_size=17, row_gap=44)

    # Sequence bar — the selling story: agents own work; humans own gates.
    rounded(d, (60, 1460, 2340, 1605), "E9FBFB", TEAL, 22, 4)
    rounded(d, (95, 1495, 165, 1565), TEAL, None, 35)
    center_text(d, (95, 1495, 165, 1565), "→", 42, WHITE, True)
    sequence = "Challenge + Research Agents  →  D1  →  Definition Agent  →  D2  →  Design Agent  →  D3  →  Validation Agent  →  D4"
    center_text(d, (185, 1475, 2310, 1590), sequence, 22, NAVY, True)
    img.save(path, quality=95)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def build_docx(path, flow_png, folder_png):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.85)
    sec.header_distance = sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, TEAL, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("AI UX DISCOVERY-TO-PROTOTYPE PIPELINE")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Workflow v2.0 • Synthetic demonstration package • 08 August 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI UX Discovery-to-Prototype Pipeline")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Workflow explanation, artifact guide, and visual architecture")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    labels = [("4", "Phases", BLUE), ("4", "Human gates", TEAL), ("22", "Reusable skills", PURPLE), ("32", "Passing tests", ORANGE)]
    for cell, (num, label, color) in zip(table.rows[0].cells, labels):
        cell.width = Inches(1.58)
        set_cell_fill(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(num + "\n")
        rr.bold = True
        rr.font.size = Pt(20)
        rr.font.color.rgb = RGBColor(255, 255, 255)
        rr = p.add_run(label)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    add_heading(doc, "What this package represents", 1)
    doc.add_paragraph("The repository demonstrates a governed AI-assisted UX process that takes a synthetic business requirements document through research, ideation, design, prototype creation, and validation. All four review gates were explicitly approved by a human reviewer.")
    add_bullet(doc, "Markdown files define phases, agents, reusable skills, gates, inputs, outputs, and evidence rules.")
    add_bullet(doc, "The runnable HTML/CSS/JavaScript prototype demonstrates the appointment-request experience.")
    add_bullet(doc, "The Python layer validates structural contracts and routes the workflow without replacing agent judgment.")
    add_bullet(doc, "The final approval applies only to the documented synthetic demo scope; it is not real-user, production, legal, or WCAG verification.")

    doc.add_page_break()
    add_heading(doc, "1. End-to-end workflow", 1)
    doc.add_paragraph("Each phase produces traceable artifacts for a human gate. APPROVE advances; REVISE returns only to affected work; REJECT stops execution. Agents cannot approve gates.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(flow_png), width=Inches(6.55))
    cap = doc.add_paragraph("Figure 1. Color workflow from BRD input to D4-approved synthetic demo prototype.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()
    add_heading(doc, "2. Phase outputs", 1)
    phase_data = [
        ("Phase 1 — Discover", BLUE, "UX Research Agent + Requirements Challenge Agent", "Challenges the BRD and synthesizes evidence without converting assumptions into user facts.", "brd-risk-review.md; persona.md; pain-points.md; research-gaps.md", "D1 — Research & Requirements Review"),
        ("Phase 2 — Define & Ideate", TEAL, "UX Definition & Ideation Agent", "Converts approved evidence into problem definitions, journeys, opportunities, and alternative concepts.", "problem-definition.md; journey-map.md; opportunities.md; concept-options.md; selected-concept.md", "D2 — Definition & Ideation Review"),
        ("Phase 3 — Design & Prototype", PURPLE, "Experience Design Agent", "Turns the selected concept into flows, screen/state specifications, and a runnable local prototype.", "user-flow.md; screen-spec.md; interaction-states.md; prototype-spec.md; prototype/", "D3 — Design & Prototype Review"),
        ("Phase 4 — Validate", ORANGE, "UX Validation & Audit Agent", "Performs synthetic walkthroughs and evidence-bounded accessibility, design-system, requirement, and edge-case checks.", "validation-report.md; validation-issues.md; requirement-coverage.md; accessibility-audit.md", "D4 — Final UX Validation Review"),
    ]
    for title, color, agents, purpose, outputs, gate in phase_data:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor.from_string(color)
        p = doc.add_paragraph()
        rr = p.add_run("Agent owner: " + agents)
        rr.bold = True
        rr.font.color.rgb = RGBColor.from_string(color)
        doc.add_paragraph(purpose)
        add_bullet(doc, "Outputs: " + outputs)
        add_bullet(doc, "Human gate: " + gate)

    add_heading(doc, "Validation outcome", 2)
    doc.add_paragraph("The core appointment-request path works, including required-field blocking, privacy consent, duplicate-submit prevention, and explicit pending/timeout wording. Accepted limitations include missing cancellation and session-expiry execution, incomplete accessible error/focus behavior, and simulated backend/notification behavior.")

    doc.add_page_break()
    add_heading(doc, "3. Folder architecture", 1)
    doc.add_paragraph("The repository separates stable workflow definitions, active project inputs/state, generated phase outputs, executable runtime code, and automated tests.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(folder_png), width=Inches(6.55))
    cap = doc.add_paragraph("Figure 2. Color-coded folder structure and responsibility boundaries.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_page_break()
    add_heading(doc, "4. Folder reference", 1)
    refs = [
        ("workflow/", "Orchestrator and four phase runbooks; controls routing and declared execution order."),
        ("agents/", "Five responsibility owners: Requirements Challenge, UX Research, UX Definition & Ideation, Experience Design, and UX Validation & Audit."),
        ("skills/", "Twenty-two reusable capabilities invoked by agents in an explicit order."),
        ("gates/", "D1–D4 human review contracts with APPROVE, REVISE, and REJECT decisions."),
        ("projects/starter/input/", "Synthetic BRD, stakeholder notes, existing flow, design system, and demo research inputs."),
        ("outputs/phase-01/", "BRD challenge, personas, pain points, and research gaps."),
        ("outputs/phase-02/", "Problem definition, journey, opportunities, concepts, and selected direction."),
        ("outputs/phase-03/", "Flow, screens, states, prototype specification, and runnable web prototype."),
        ("outputs/phase-04/", "Validation report, issues, requirement coverage, and accessibility audit."),
        ("outputs/reviews/", "Explicit human decision records for D1 through D4."),
        ("src/ai_ux_workflow/", "Python validation, routing, command-line interface, and deterministic challenge demo."),
        ("tests/", "Contract, cross-reference, safety, sample-input, assignment-compliance, and runtime tests."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    hdr = table.rows[0].cells
    hdr[0].text = "Folder / file"
    hdr[1].text = "Responsibility"
    for c in hdr:
        set_cell_fill(c, NAVY)
        for r in c.paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
    for idx, (name, desc) in enumerate(refs):
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = desc
        if idx % 2:
            set_cell_fill(cells[0], "F1F5F9")
            set_cell_fill(cells[1], "F1F5F9")
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        for cell in cells:
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_heading(doc, "Usage note", 2)
    doc.add_paragraph("The Markdown workflow can be used in compatible agentic environments by instructing the agent to read the orchestrator, active project configuration/state, and the relevant phase runbook. The Python runtime is an optional deterministic validation and demonstration layer, not the only way to use the workflow.")
    doc.save(path)


if __name__ == "__main__":
    flow = OUT / "AI_UX_Workflow_Flow_Color.png"
    folders = OUT / "AI_UX_Folder_Structure_Color.png"
    docx = OUT / "AI_UX_Workflow_Visual_Guide_Agent_Tree.docx"
    workflow_image(flow)
    folder_image(folders)
    build_docx(docx, flow, folders)
    print(flow)
    print(folders)
    print(docx)
